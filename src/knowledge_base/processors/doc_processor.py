"""
DOC/DOCX文件处理器
专门用于处理基层治理案例的Word文档
"""
import os
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
from langchain_core.documents import Document
from src.utils.logger import logger

class DocProcessor:
    """DOC/DOCX文件处理器"""
    
    def __init__(self):
        """初始化处理器"""
        self.supported_formats = ['.doc', '.docx']
        
        # 案例结构关键词
        self.case_keywords = {
            'title': ['标题', '案例名称', '案例标题', '题目'],
            'problem': ['问题', '问题描述', '背景', '情况', '现状'],
            'solution': ['解决方案', '处理过程', '解决步骤', '措施', '做法'],
            'result': ['结果', '效果', '成果', 'outcome'],
            'reflection': ['经验', '启示', '总结', '反思', '体会']
        }
    
    def process_doc_files(self, directory: str) -> List[Document]:
        """
        处理目录下的所有DOC/DOCX文件
        
        Args:
            directory: 文件目录路径
            
        Returns:
            处理后的文档列表
        """
        logger.info(f"开始处理目录: {directory}")
        
        documents = []
        processed_count = 0
        failed_count = 0
        
        for root, dirs, files in os.walk(directory):
            for file in files:
                if any(file.lower().endswith(fmt) for fmt in self.supported_formats):
                    file_path = os.path.join(root, file)
                    
                    try:
                        doc = self._process_single_file(file_path)
                        if doc:
                            documents.append(doc)
                            processed_count += 1
                        else:
                            failed_count += 1
                    except Exception as e:
                        logger.warning(f"处理文件失败 {file_path}: {e}")
                        failed_count += 1
        
        logger.info(f"文档处理完成: 成功 {processed_count} 个, 失败 {failed_count} 个")
        return documents
    
    def _process_single_file(self, file_path: str) -> Optional[Document]:
        """
        处理单个文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            处理后的文档对象
        """
        try:
            # 提取文本内容
            content = self._extract_text_content(file_path)
            
            if not content or len(content.strip()) < 50:
                logger.warning(f"文件内容过短或为空: {file_path}")
                return None
            
            # 解析案例结构
            structured_content = self._parse_case_structure(content)
            
            # 提取元数据
            metadata = self._extract_metadata(file_path, structured_content)
            
            # 格式化内容
            formatted_content = self._format_case_content(structured_content)
            
            return Document(
                page_content=formatted_content,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"处理文件失败 {file_path}: {e}")
            return None
    
    def _extract_text_content(self, file_path: str) -> str:
        """提取文本内容"""
        try:
            if file_path.lower().endswith('.docx'):
                return self._extract_docx_content(file_path)
            elif file_path.lower().endswith('.doc'):
                return self._extract_doc_content(file_path)
            else:
                return ""
        except Exception as e:
            logger.error(f"提取文本内容失败 {file_path}: {e}")
            return ""
    
    def _extract_docx_content(self, file_path: str) -> str:
        """提取DOCX文件内容"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
            
            # 处理表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        content_parts.append(' | '.join(row_text))
            
            return '\n'.join(content_parts)
            
        except ImportError:
            logger.error("需要安装python-docx库: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"提取DOCX内容失败: {e}")
            return ""
    
    def _extract_doc_content(self, file_path: str) -> str:
        """提取DOC文件内容"""
        try:
            # 尝试使用python-docx2txt
            try:
                import docx2txt
                content = docx2txt.process(file_path)
                return content if content else ""
            except ImportError:
                pass
            
            # 尝试使用win32com (仅Windows)
            try:
                import win32com.client
                
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                doc = word.Documents.Open(file_path)
                content = doc.Content.Text
                doc.Close()
                word.Quit()
                
                return content
            except ImportError:
                pass
            
            # 如果都不可用，记录警告
            logger.warning(f"无法处理.doc文件 {file_path}，请安装docx2txt或在Windows环境下运行")
            return ""
            
        except Exception as e:
            logger.error(f"提取DOC内容失败: {e}")
            return ""
    
    def _parse_case_structure(self, content: str) -> Dict[str, str]:
        """解析案例结构"""
        structured = {
            'title': '',
            'problem': '',
            'solution': '',
            'result': '',
            'reflection': '',
            'raw_content': content
        }
        
        # 按段落分割内容
        paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
        
        # 尝试识别结构化内容
        current_section = None
        current_content = []
        
        for paragraph in paragraphs:
            # 检查是否是标题行
            section_found = False
            for section, keywords in self.case_keywords.items():
                for keyword in keywords:
                    if keyword in paragraph and len(paragraph) < 50:
                        # 保存之前的内容
                        if current_section and current_content:
                            structured[current_section] = '\n'.join(current_content)
                        
                        # 开始新的部分
                        current_section = section
                        current_content = []
                        section_found = True
                        break
                if section_found:
                    break
            
            if not section_found and current_section:
                current_content.append(paragraph)
            elif not current_section and not structured['title']:
                # 第一个段落可能是标题
                structured['title'] = paragraph
        
        # 保存最后一个部分
        if current_section and current_content:
            structured[current_section] = '\n'.join(current_content)
        
        # 如果没有识别出标题，使用文件名
        if not structured['title']:
            structured['title'] = '未命名案例'
        
        return structured
    
    def _extract_metadata(self, file_path: str, structured_content: Dict[str, str]) -> Dict[str, Any]:
        """提取元数据"""
        metadata = {
            'source': file_path,
            'filename': Path(file_path).name,
            'title': structured_content.get('title', Path(file_path).stem),
            'category': self._infer_category(structured_content),
            'file_type': Path(file_path).suffix,
            'has_structure': self._check_structure_completeness(structured_content)
        }
        
        # 从文件路径推断额外信息
        path_parts = Path(file_path).parts
        if len(path_parts) > 2:
            # 可能包含作者或分类信息
            parent_dir = path_parts[-2]
            metadata['author_or_category'] = parent_dir
        
        return metadata
    
    def _infer_category(self, structured_content: Dict[str, str]) -> str:
        """推断案例类别"""
        # 合并所有文本内容进行分析
        all_text = ' '.join([
            structured_content.get('title', ''),
            structured_content.get('problem', ''),
            structured_content.get('solution', ''),
            structured_content.get('raw_content', '')
        ]).lower()
        
        category_keywords = {
            '邻里纠纷': ['邻里', '纠纷', '矛盾', '争吵', '冲突', '邻居'],
            '民生服务': ['救助', '帮扶', '困难', '低保', '医疗', '养老', '就业'],
            '社区治理': ['停车', '环境', '卫生', '管理', '秩序', '物业', '小区'],
            '政策宣传': ['宣传', '政策', '解读', '培训', '教育', '普及'],
            '环境治理': ['环境', '垃圾', '绿化', '污染', '整治', '清洁'],
            '安全管理': ['安全', '消防', '治安', '防范', '巡逻', '监控'],
            '文化活动': ['文化', '活动', '娱乐', '节庆', '表演', '比赛'],
            '经济发展': ['经济', '发展', '产业', '就业', '创业', '扶贫'],
            '教育服务': ['教育', '学校', '学生', '培训', '学习'],
            '医疗卫生': ['医疗', '卫生', '健康', '疫情', '防控']
        }
        
        # 计算每个类别的匹配分数
        category_scores = {}
        for category, keywords in category_keywords.items():
            score = sum(1 for keyword in keywords if keyword in all_text)
            if score > 0:
                category_scores[category] = score
        
        # 返回得分最高的类别
        if category_scores:
            return max(category_scores.items(), key=lambda x: x[1])[0]
        else:
            return '其他'
    
    def _check_structure_completeness(self, structured_content: Dict[str, str]) -> float:
        """检查结构完整性"""
        required_fields = ['title', 'problem', 'solution']
        present_fields = sum(1 for field in required_fields 
                           if structured_content.get(field, '').strip())
        
        return present_fields / len(required_fields)
    
    def _format_case_content(self, structured_content: Dict[str, str]) -> str:
        """格式化案例内容"""
        formatted_parts = []
        
        # 标题
        if structured_content.get('title'):
            formatted_parts.append(f"案例标题：{structured_content['title']}")
        
        # 问题描述
        if structured_content.get('problem'):
            formatted_parts.append(f"问题描述：\n{structured_content['problem']}")
        
        # 解决方案
        if structured_content.get('solution'):
            formatted_parts.append(f"解决方案：\n{structured_content['solution']}")
        
        # 结果
        if structured_content.get('result'):
            formatted_parts.append(f"处理结果：\n{structured_content['result']}")
        
        # 经验总结
        if structured_content.get('reflection'):
            formatted_parts.append(f"经验总结：\n{structured_content['reflection']}")
        
        # 如果结构化内容不足，添加原始内容
        if len(formatted_parts) < 3:
            formatted_parts.append(f"原始内容：\n{structured_content['raw_content']}")
        
        return '\n\n'.join(formatted_parts)
    
    def get_processing_statistics(self, directory: str) -> Dict[str, Any]:
        """获取处理统计信息"""
        stats = {
            'total_files': 0,
            'docx_files': 0,
            'doc_files': 0,
            'file_sizes': [],
            'categories': {},
            'authors': set()
        }
        
        for root, dirs, files in os.walk(directory):
            for file in files:
                if any(file.lower().endswith(fmt) for fmt in self.supported_formats):
                    file_path = os.path.join(root, file)
                    stats['total_files'] += 1
                    
                    if file.lower().endswith('.docx'):
                        stats['docx_files'] += 1
                    elif file.lower().endswith('.doc'):
                        stats['doc_files'] += 1
                    
                    # 文件大小
                    try:
                        size = os.path.getsize(file_path)
                        stats['file_sizes'].append(size)
                    except OSError:
                        pass
                    
                    # 从路径推断作者
                    path_parts = Path(file_path).parts
                    if len(path_parts) > 2:
                        author = path_parts[-2]
                        stats['authors'].add(author)
        
        # 计算统计信息
        if stats['file_sizes']:
            stats['avg_file_size'] = sum(stats['file_sizes']) / len(stats['file_sizes'])
            stats['total_size'] = sum(stats['file_sizes'])
        else:
            stats['avg_file_size'] = 0
            stats['total_size'] = 0
        
        stats['authors'] = list(stats['authors'])
        
        return stats

if __name__ == "__main__":
    # 测试文档处理器
    processor = DocProcessor()
    
    case_dir = "./data/knowledge_base/已有案例"
    
    if os.path.exists(case_dir):
        # 获取统计信息
        stats = processor.get_processing_statistics(case_dir)
        print("文档统计信息:")
        print(f"  总文件数: {stats['total_files']}")
        print(f"  DOCX文件: {stats['docx_files']}")
        print(f"  DOC文件: {stats['doc_files']}")
        print(f"  总大小: {stats['total_size'] / 1024 / 1024:.2f} MB")
        print(f"  作者/分类: {stats['authors']}")
        
        # 处理文档
        print("\n开始处理文档...")
        documents = processor.process_doc_files(case_dir)
        
        print(f"\n处理结果:")
        print(f"  成功处理: {len(documents)} 个文档")
        
        # 显示前几个案例
        for i, doc in enumerate(documents[:3], 1):
            print(f"\n案例 {i}:")
            print(f"  标题: {doc.metadata.get('title', '未知')}")
            print(f"  类别: {doc.metadata.get('category', '未知')}")
            print(f"  内容预览: {doc.page_content[:200]}...")
    else:
        print(f"案例目录不存在: {case_dir}")