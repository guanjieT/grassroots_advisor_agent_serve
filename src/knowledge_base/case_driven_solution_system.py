"""
案例驱动的基层治理解决方案生成系统
以成功治理案例为模板，结合当地法规生成定制化解决方案
"""
import os
import json
import re
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from langchain_core.documents import Document
from langchain_community.embeddings import DashScopeEmbeddings
from langchain_chroma import Chroma
from langchain_community.llms import Tongyi
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from config import config
from utils.logger import logger

class CaseDrivenSolutionSystem:
    """案例驱动的解决方案生成系统"""
    
    def __init__(self):
        """初始化系统"""
        logger.info("初始化案例驱动解决方案系统...")
        
        # 初始化嵌入模型
        self.embeddings = DashScopeEmbeddings(
            dashscope_api_key=config.DASHSCOPE_API_KEY,
            model=config.EMBEDDING_MODEL
        )
        
        # 初始化LLM
        self.llm = Tongyi(
            dashscope_api_key=config.DASHSCOPE_API_KEY,
            model_name=config.LLM_MODEL,
            temperature=0.3,  # 稍微降低温度，确保方案的一致性
            max_tokens=2000
        )
        
        # 案例向量数据库
        self.case_vectorstore = None
        self.policy_vectorstore = None
        
        # 初始化向量数据库
        self._initialize_vectorstores()
        
        # 设置提示模板
        self._setup_prompts()
        
        logger.info("案例驱动解决方案系统初始化完成")
    
    def _initialize_vectorstores(self):
        """初始化向量数据库"""
        try:
            # 案例向量数据库
            case_persist_dir = "./data/case_vectorstore"
            os.makedirs(case_persist_dir, exist_ok=True)
            
            self.case_vectorstore = Chroma(
                collection_name="governance_cases",
                embedding_function=self.embeddings,
                persist_directory=case_persist_dir
            )
            
            # 政策向量数据库（复用已有的）
            policy_persist_dir = "./data/policy_rag_chroma"
            if os.path.exists(policy_persist_dir):
                self.policy_vectorstore = Chroma(
                    collection_name="policy_knowledge",
                    embedding_function=self.embeddings,
                    persist_directory=policy_persist_dir
                )
            
            logger.info("向量数据库初始化完成")
            
        except Exception as e:
            logger.error(f"向量数据库初始化失败: {e}")
            raise
    
    def _setup_prompts(self):
        """设置提示模板"""
        
        # 解决方案生成提示模板
        self.solution_prompt = PromptTemplate(
            input_variables=["problem", "location", "similar_cases", "relevant_policies"],
            template="""你是一个专业的基层治理顾问，请基于成功案例和相关政策，为特定地区的基层治理问题生成具体可执行的解决方案。

## 问题描述
**地区**: {location}
**问题**: {problem}

## 相似成功案例
{similar_cases}

## 相关政策法规
{relevant_policies}

## 请生成解决方案，包含以下内容：

### 1. 问题分析
- 问题的根本原因分析
- 涉及的利益相关方
- 可能的风险点

### 2. 解决方案设计
- **总体思路**: 基于成功案例的核心经验
- **具体步骤**: 详细的执行步骤（至少5步）
- **时间安排**: 每个步骤的预期时间
- **责任分工**: 明确各方职责

### 3. 政策合规性
- 相关法律法规依据
- 合规性要求和注意事项
- 需要的审批或备案程序

### 4. 实施保障
- 所需资源和条件
- 风险防控措施
- 效果评估标准

### 5. 地区适配
- 结合{location}的具体情况
- 考虑当地的文化和习俗
- 可能需要的本地化调整

### 6. 预期效果
- 短期目标（1-3个月）
- 中期目标（3-6个月）
- 长期目标（6个月以上）

请确保方案具体可操作，符合当地法规要求，并充分借鉴成功案例的经验。

解决方案："""
        )
        
        # 案例相似度分析提示模板
        self.similarity_prompt = PromptTemplate(
            input_variables=["problem", "case_content"],
            template="""请分析以下问题与案例的相似度，并提取可借鉴的关键经验。

问题：{problem}

案例：{case_content}

请从以下角度分析：
1. 问题类型的相似度（0-1分）
2. 解决思路的可借鉴性
3. 关键成功因素
4. 可直接应用的具体措施
5. 需要调整的地方

分析结果："""
        )
    
    def build_case_knowledge_base(self, case_dir: str = "./data/knowledge_base/已有案例") -> bool:
        """
        构建案例知识库
        
        Args:
            case_dir: 案例文件目录
            
        Returns:
            是否构建成功
        """
        logger.info("开始构建案例知识库...")
        
        try:
            # 清空现有数据
            self._clear_case_vectorstore()
            
            # 加载案例文档
            case_documents = self._load_case_documents(case_dir)
            
            if not case_documents:
                logger.warning("没有找到案例文档")
                return False
            
            # 添加到向量数据库
            batch_size = 50
            for i in range(0, len(case_documents), batch_size):
                batch = case_documents[i:i + batch_size]
                self.case_vectorstore.add_documents(batch)
                logger.info(f"已处理 {min(i + batch_size, len(case_documents))}/{len(case_documents)} 个案例")
            
            logger.info(f"案例知识库构建完成，共处理 {len(case_documents)} 个案例")
            return True
            
        except Exception as e:
            logger.error(f"构建案例知识库失败: {e}")
            return False
    
    def _load_case_documents(self, case_dir: str) -> List[Document]:
        """加载案例文档"""
        documents = []
        
        # 1. 加载JSON格式的样本案例
        json_file = os.path.join("./data/knowledge_base", "sample_cases.json")
        if os.path.exists(json_file):
            with open(json_file, 'r', encoding='utf-8') as f:
                sample_cases = json.load(f)
            
            for case in sample_cases:
                content = self._format_case_content(case)
                doc = Document(
                    page_content=content,
                    metadata={
                        'case_id': case['id'],
                        'title': case['title'],
                        'category': case['category'],
                        'source': 'sample_cases',
                        'keywords': ','.join(case.get('keywords', []))
                    }
                )
                documents.append(doc)
        
        # 2. 加载DOC/DOCX格式的案例文件
        if os.path.exists(case_dir):
            for root, dirs, files in os.walk(case_dir):
                for file in files:
                    if file.endswith(('.doc', '.docx')):
                        file_path = os.path.join(root, file)
                        try:
                            content = self._extract_doc_content(file_path)
                            if content and len(content.strip()) > 100:
                                doc = Document(
                                    page_content=content,
                                    metadata={
                                        'title': Path(file).stem,
                                        'source': file_path,
                                        'category': self._infer_category_from_content(content),
                                        'file_type': Path(file).suffix
                                    }
                                )
                                documents.append(doc)
                        except Exception as e:
                            logger.warning(f"处理文件失败 {file_path}: {e}")
                            continue
        
        return documents
    
    def _format_case_content(self, case: Dict[str, Any]) -> str:
        """格式化案例内容"""
        content = f"""案例标题：{case['title']}
案例类别：{case['category']}

问题描述：
{case['problem']}

解决步骤：
"""
        for i, step in enumerate(case['steps'], 1):
            content += f"{i}. {step}\n"
        
        content += f"""
解决结果：
{case['result']}

经验总结：
{case['reflection']}

关键词：{', '.join(case.get('keywords', []))}
"""
        return content
    
    def _extract_doc_content(self, file_path: str) -> str:
        """提取DOC/DOCX文件内容"""
        try:
            if file_path.endswith('.docx'):
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text.strip())
                return '\n'.join(content)
            else:
                # 对于.doc文件，可以使用python-docx2txt或其他库
                # 这里先返回空，实际使用时需要安装相应的库
                logger.warning(f"暂不支持.doc格式文件: {file_path}")
                return ""
        except Exception as e:
            logger.error(f"提取文档内容失败 {file_path}: {e}")
            return ""
    
    def _infer_category_from_content(self, content: str) -> str:
        """从内容推断案例类别"""
        category_keywords = {
            '邻里纠纷': ['邻里', '纠纷', '矛盾', '争吵', '冲突'],
            '民生服务': ['救助', '帮扶', '困难', '低保', '医疗'],
            '社区治理': ['停车', '环境', '卫生', '管理', '秩序'],
            '政策宣传': ['宣传', '政策', '解读', '培训', '教育'],
            '环境治理': ['环境', '垃圾', '绿化', '污染', '整治'],
            '安全管理': ['安全', '消防', '治安', '防范', '巡逻'],
            '文化活动': ['文化', '活动', '娱乐', '节庆', '表演']
        }
        
        for category, keywords in category_keywords.items():
            if any(keyword in content for keyword in keywords):
                return category
        
        return '其他'
    
    def _clear_case_vectorstore(self):
        """清空案例向量数据库"""
        try:
            self.case_vectorstore.delete_collection()
            self._initialize_vectorstores()
            logger.info("案例向量数据库已清空")
        except Exception as e:
            logger.warning(f"清空案例向量数据库失败: {e}")
    
    def generate_solution(
        self, 
        problem: str, 
        location: str = "某地区",
        case_count: int = 3,
        policy_count: int = 3
    ) -> Dict[str, Any]:
        """
        生成解决方案
        
        Args:
            problem: 问题描述
            location: 地区名称
            case_count: 参考案例数量
            policy_count: 参考政策数量
            
        Returns:
            解决方案信息
        """
        logger.info(f"为{location}的问题生成解决方案: {problem}")
        
        try:
            # 1. 检索相似案例
            similar_cases = self._find_similar_cases(problem, case_count)
            
            # 2. 检索相关政策
            relevant_policies = self._find_relevant_policies(problem, policy_count)
            
            # 3. 生成解决方案
            solution = self._generate_solution_content(
                problem, location, similar_cases, relevant_policies
            )
            
            return {
                'problem': problem,
                'location': location,
                'solution': solution,
                'reference_cases': [
                    {
                        'title': case.metadata.get('title', '未知'),
                        'category': case.metadata.get('category', '未知'),
                        'content_preview': case.page_content[:200] + "..."
                    }
                    for case in similar_cases
                ],
                'reference_policies': [
                    {
                        'title': policy.metadata.get('title', '未知'),
                        'admin_level': policy.metadata.get('admin_level', '未知'),
                        'content_preview': policy.page_content[:200] + "..."
                    }
                    for policy in relevant_policies
                ],
                'generation_time': self._get_current_time()
            }
            
        except Exception as e:
            logger.error(f"生成解决方案失败: {e}")
            return {
                'problem': problem,
                'location': location,
                'solution': f'生成解决方案时出现错误: {str(e)}',
                'reference_cases': [],
                'reference_policies': [],
                'generation_time': self._get_current_time()
            }
    
    def _find_similar_cases(self, problem: str, k: int = 3) -> List[Document]:
        """查找相似案例"""
        if not self.case_vectorstore:
            logger.warning("案例向量数据库未初始化")
            return []
        
        try:
            cases = self.case_vectorstore.similarity_search(problem, k=k)
            logger.info(f"找到 {len(cases)} 个相似案例")
            return cases
        except Exception as e:
            logger.error(f"检索相似案例失败: {e}")
            return []
    
    def _find_relevant_policies(self, problem: str, k: int = 3) -> List[Document]:
        """查找相关政策"""
        if not self.policy_vectorstore:
            logger.warning("政策向量数据库未初始化")
            return []
        
        try:
            policies = self.policy_vectorstore.similarity_search(problem, k=k)
            logger.info(f"找到 {len(policies)} 个相关政策")
            return policies
        except Exception as e:
            logger.error(f"检索相关政策失败: {e}")
            return []
    
    def _generate_solution_content(
        self, 
        problem: str, 
        location: str, 
        cases: List[Document], 
        policies: List[Document]
    ) -> str:
        """生成解决方案内容"""
        
        # 格式化案例内容
        cases_text = ""
        for i, case in enumerate(cases, 1):
            cases_text += f"""
### 案例 {i}: {case.metadata.get('title', '未知')}
**类别**: {case.metadata.get('category', '未知')}
**内容**: {case.page_content}
---
"""
        
        # 格式化政策内容
        policies_text = ""
        for i, policy in enumerate(policies, 1):
            policies_text += f"""
### 政策 {i}: {policy.metadata.get('title', '未知')}
**层级**: {policy.metadata.get('admin_level', '未知')}
**内容**: {policy.page_content}
---
"""
        
        # 生成解决方案
        try:
            solution = self.solution_prompt.format(
                problem=problem,
                location=location,
                similar_cases=cases_text,
                relevant_policies=policies_text
            )
            
            result = self.llm.invoke(solution)
            return result
            
        except Exception as e:
            logger.error(f"生成解决方案内容失败: {e}")
            return f"生成解决方案时出现错误: {str(e)}"
    
    def analyze_case_similarity(self, problem: str, case_content: str) -> str:
        """分析案例相似度"""
        try:
            analysis = self.similarity_prompt.format(
                problem=problem,
                case_content=case_content
            )
            
            result = self.llm.invoke(analysis)
            return result
            
        except Exception as e:
            logger.error(f"分析案例相似度失败: {e}")
            return f"分析失败: {str(e)}"
    
    def get_case_statistics(self) -> Dict[str, Any]:
        """获取案例库统计信息"""
        try:
            if not self.case_vectorstore:
                return {'total_cases': 0, 'categories': {}}
            
            # 这里需要实现统计逻辑
            # ChromaDB没有直接的统计接口，需要通过查询实现
            return {
                'total_cases': 'N/A',
                'categories': 'N/A',
                'note': '需要实现统计功能'
            }
            
        except Exception as e:
            logger.error(f"获取案例统计失败: {e}")
            return {'error': str(e)}
    
    def _get_current_time(self) -> str:
        """获取当前时间"""
        from datetime import datetime
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def build_case_driven_system() -> CaseDrivenSolutionSystem:
    """构建案例驱动系统的便捷函数"""
    logger.info("构建案例驱动解决方案系统...")
    
    try:
        # 创建系统实例
        system = CaseDrivenSolutionSystem()
        
        # 构建案例知识库
        success = system.build_case_knowledge_base()
        
        if success:
            logger.info("案例驱动解决方案系统构建成功")
            return system
        else:
            logger.error("案例知识库构建失败")
            return None
            
    except Exception as e:
        logger.error(f"构建案例驱动系统失败: {e}")
        return None

if __name__ == "__main__":
    # 构建和测试系统
    system = build_case_driven_system()
    
    if system:
        # 测试解决方案生成
        test_problems = [
            {
                'problem': '小区内经常有居民遛狗不拴绳，其他居民感到害怕，多次投诉无果',
                'location': '北京市朝阳区某社区'
            },
            {
                'problem': '老年人对智能手机操作不熟悉，无法使用健康码，影响出行',
                'location': '上海市浦东新区某街道'
            },
            {
                'problem': '小区垃圾分类执行不到位，居民积极性不高',
                'location': '广州市天河区某小区'
            }
        ]
        
        for i, test_case in enumerate(test_problems, 1):
            print(f"\n{'='*60}")
            print(f"测试案例 {i}")
            print(f"{'='*60}")
            
            result = system.generate_solution(
                test_case['problem'], 
                test_case['location']
            )
            
            print(f"问题: {result['problem']}")
            print(f"地区: {result['location']}")
            print(f"\n参考案例 ({len(result['reference_cases'])} 个):")
            for j, case in enumerate(result['reference_cases'], 1):
                print(f"  {j}. {case['title']} ({case['category']})")
            
            print(f"\n参考政策 ({len(result['reference_policies'])} 个):")
            for j, policy in enumerate(result['reference_policies'], 1):
                print(f"  {j}. {policy['title']} ({policy['admin_level']})")
            
            print(f"\n解决方案:")
            print(result['solution'][:500] + "..." if len(result['solution']) > 500 else result['solution'])
    else:
        print("系统构建失败")